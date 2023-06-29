# -*- coding: utf-8 -*-
"""
Created on Tue Jun 27 11:42:04 2023

@author: 86136
"""

"""
实现功能：文档解析
输入：Microsoft Word
输出：文档目录、文档内图片（PNG格式）、文档表格（采用pandas的dataFrame数据类型表示）、可返回的文档版面数据
"""

import argparse
import docx
import pandas as pd
import os, re

from docx import Document


def docx_extract_tables(doc_path):
    """
    功能：从Microsoft Word中提取表格
    输入：Microsoft Word文件路径doc_path（"example.docx"）
    输出：对于每一个表输出一个pandas的dataFrame类型，采用print方式
    返回值：表格df
    """
    
    # open the document and new a document object to handle it
    document = docx.Document(doc_path)
    result = []
    
    #  the table to a list
    for tb in document.tables:
        for i, row in enumerate(tb.rows):
            temp = []
            for j, cell in enumerate(row.cells):
                text = ''
                for p in cell.paragraphs:
                    text += p.text
                temp.append(text)
            result.append(temp)
        
        # use a dataFrame object to handle the result
        df = pd.DataFrame(result[1:], columns=result[0])
        print(df)
    return df
        

def docx_extract_images(word_path, result_path):
    """
    功能：从Microsoft Word中提取图片（从docx文件中提取所有图片，并将它们保存为PNG文件）
    输入：Microsoft Word文件路径doc_path（"example.docx"）
    输出：对于每一个图片以png格式直接输出到指定路径result_path
    返回值：无
    """
    
    # create a document object to handle the doc file
    doc = docx.Document(word_path)
    
    # extract the rels
    dict_rel = doc.part._rels
    for rel in dict_rel:
        rel = dict_rel[rel]
        if "image" in rel.target_ref:
            
            # make sure the result path exists
            if not os.path.exists(result_path):
                os.makedirs(result_path)
                
            # spawn the image name
            img_name = re.findall("/(.*)", rel.target_ref)[0]
            word_name = os.path.splitext(word_path)[0]
            if os.sep in word_name:
                new_name = word_name.split('\\')[-1]
            else:
                new_name = word_name.split('/')[-1]
            img_name = f'{new_name}-'+f'{img_name}'
            
            # create the image file and write in the content
            with open(f'{result_path}/{img_name}', "wb") as f:
                f.write(rel.target_part.blob)
                

def docx_extract_toc(doc_path):
    """
    功能：从Microsoft Word中提取目录并输出
    输入：Microsoft Word文件路径doc_path（"example.docx"）
    输出：一个list对象，包括每个章节的标题
    返回值：章节的标题list
    """
    
    # create a document object to handle the doc
    document = Document(doc_path)
    result = []
    
    # extract the toc with the label 'Heading'
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            result.append(paragraph.text)
    print(result)
    return result


def docx_extract_text(doc_path):
    """
    功能：从Microsoft Word中提取正文并输出
    输入：Microsoft Word文件路径doc_path（"example.docx"）
    输出：输出正文内容
    返回值：正文list
    """
    
    # extract the plaintext
    document = Document(doc_path)
    result = []
    for paragraph in document.paragraphs:
        result.append(paragraph.text)
    print(result)
    return result
    

if __name__ == "__main__":
    # add a parser
    parser = argparse.ArgumentParser(description='Microsoft Word Extraction')
    
    # add cmd args
    parser.add_argument('-o','--option', type=str, default='toc', help='The function will be served, optional choice will be text(Extract text), image(Extract images), table(Extract tables) and toc(Extract tocs)')
    parser.add_argument('-p','--path', type=str, default='demo.docx', help='Path to the doc file')
    parser.add_argument('-r','--result_path', type=str, default='./', help='Path to store the result images')
    
    # parse the args
    args = parser.parse_args()
    
    # parse the options
    if args.option == 'text':
        docx_extract_text(args.path)
    elif args.option == 'image':
        docx_extract_images(args.path, args.result_path)
    elif args.option == 'table':
        docx_extract_tables(args.path)
    elif args.option == 'toc':
        docx_extract_toc(args.path)
    else:
        print("Invalid options.")